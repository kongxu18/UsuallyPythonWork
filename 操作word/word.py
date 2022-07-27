"""
生成报告
"""
from docx import Document
from docx.shared import Inches
from settings import DATA, LOC

document = Document('source/TestReport.docx')

print(document)

tables = document.tables

table = tables[2]
for i, row in enumerate(table.rows):  # 读每行
    for j, cell in enumerate(row.cells):  # 读一行中的所有单元格
        c = cell.text

        break


def write(table, name, loc):
    data = DATA[name]
    if name == 'cover_2':
        for key, val in data.items():
            row, col = loc[key]
            paragraphs = table.rows[row].cells[col].paragraphs
            paragraphs[0].text = val[0]
            try:
                paragraphs[1].text = val[1]
            except IndexError as err:
                pass
        return

    for key, val in data.items():
        row, col = loc[key]
        table.rows[row].cells[col].text = val
        # print(key, val, row, col)


# 封面 1 部分
cover_1 = tables[1]
write(cover_1, 'cover_1', LOC)

# 封面 2 部分
cover_2 = tables[2]
write(cover_2, 'cover_2', LOC)

# 实验数据表格

document.save('demo.docx')
