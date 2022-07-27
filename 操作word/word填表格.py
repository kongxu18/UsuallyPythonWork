import pymssql as mssql
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


class Connection(object):
    def __init__(self):
        self.sql = "select * from T257C大连索网 cross apply FT254E索网质检记录(全局代码,轴号)"

    def connect(self):
        df = None
        try:
            # sql服务器名，这里(127.0.0.1)是本地数据库IP
            serverName = 'erp.highbird.cn:9155'
            # 登陆用户名和密码
            userName = 'nizihua'
            passWord = '333333'
            # 建立连接并获取cursor
            conn = mssql.connect(serverName, userName, passWord, "base1", 'utf8')
            cursor = conn.cursor()
            cursor.execute(self.sql)
            # 列名

            columns = [name[0] for name in cursor.description]

            df = pd.DataFrame(cursor.fetchall(), columns=columns)

            cursor.close()
            conn.close()
        except Exception as err:
            print(err, 'sql---错误')

        return df

    @property
    def data(self):
        """
        字段处理
        :return:
        """
        df = self.connect()
        print(df)
        dataframe = df.groupby(by=['全局代码', '轴号'])
        return dataframe


class Word(object):
    """
    操作word
    """

    def __init__(self, src_path, target_path):
        self.src_path = src_path
        self.target_path = target_path
        self.month = '5'
        self.day = '20'

    def initialize(self, argument, df):
        document = Document(self.src_path)

        table = document.tables[0]

        # 先处理首行不规则的单元格
        level, axis = argument
        cell = table.cell(0, 2)
        p = cell.paragraphs[0]

        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run('    %s    ' % level)
        run.font.underline = True
        p.add_run('层')
        run = p.add_run('    %s    ' % axis)
        run.font.underline = True
        p.add_run('轴')

        # cell.text = '    %s     层     %s    轴' % (level, axis)

        # cell = table.cell(0, 8)
        # p = cell.paragraphs[0]
        # p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # p.add_run('2022年')
        # run = p.add_run('   %s   ' % self.month)
        # run.font.underline = True
        # p.add_run('月')
        # run = p.add_run('   %s   ' % self.day)
        # run.font.underline = True
        # p.add_run('日')

        df_rows = df.shape[0]
        df_cols = df.shape[1]
        height = table.rows[2].height
        print('row', df_rows)
        if df_rows > 10:
            # 表格增加行数
            for i in range(11, df_rows + 2):
                row = table.add_row()
                row.height = height
                if i == df_rows + 1:
                    # 这是最后一行
                    self.add(row, 0, 1, '施工')
                    self.add(row, 2, 3, None)
                    self.add(row, 4, 5, '检查')
                    self.add(row, 6, 7, None)
                    self.add(row, 8, 8, '审核')
                    self.add(row, 9, 10, None)
                else:
                    cell = row.cells[0]
                    # 单元格居中
                    cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p = cell.paragraphs[0]
                    # 段落居中
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run(str(i))
                    run.font.name = u'Calibri (正文)'
                    run.font.size = Pt(12)
        # 特殊处理前十行序号保持一致
        for i in range(1, 11):
            cell = table.cell(1 + i, 0)
            # 单元格居中
            cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            # 段落居中
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(str(i))
            run.font.name = u'Calibri (正文)'
            run.font.size = Pt(12)

        # 遍历df
        df = df.reset_index(drop=True)
        print(df)
        for index, row in df.iterrows():
            self.write(table, index, row)

        # for i, row in enumerate(table.rows):  # 读每行
        #     if i == 1:
        #         for j, cell in enumerate(row.cells):  # 读一行中的所有单元格
        #             c = cell.text
        #
        #             print(c, 'col')

        document.save('%s-%s-钢索张拉标定质量检查记录.docx' % (level, axis))

    def write(self, table, i, row):
        序号 = str(i + 1)
        钢索编号 = row['钢索编号']
        设计直径 = row['设计直径']
        设计张力 = row['设计张力']
        设计标定点位数 = row['设计标定点位数']
        设计索长 = row['设计索长']
        实测张力 = row['实测张力']
        施工标定点位数 = row['施工标定点位数']
        实测索长 = row['实测索长']
        索长偏差 = row['索长偏差']
        合格判定 = row['合格判定']
        li = [序号, 钢索编号, 设计直径, 设计张力, 设计标定点位数, 设计索长, 实测张力, 施工标定点位数, 实测索长, 索长偏差, 合格判定]
        # table 从第2行开始写,一共11列，从1开始写
        for col in range(1, 11):
            cell = table.cell(2 + i, col)
            # 单元格居中
            cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            # 段落居中
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(str(li[col]))

    # 添加最后一行
    def add(self, row_obj, start, end, word):
        # 这是最后一行
        cell_1 = row_obj.cells[start]
        cell_2 = row_obj.cells[end]
        if word:
            cell_1.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p = cell_1.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(word)
            run.font.name = u'Calibri (正文)'
            run.font.size = Pt(12)
        cell_1.merge(cell_2)


if __name__ == '__main__':
    c = Connection().data

    w = Word('钢索张拉标定质量检查记录 220516.docx', '操作word/钢索张拉word/')
    # w.initialize(('1','2'),'w')
    for argument, data in c:
        w.initialize(argument, data)

